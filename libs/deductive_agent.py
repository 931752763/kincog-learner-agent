"""
deductive_agent.py

演绎Agent：导入Word文档，生成大纲，分步讲解，支持RAG问答与动态调整。
"""
from typing import List, Any
from docx import Document
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import DashScopeEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LCDocument
from agent.llm import llm  # 你的 LLM 实例

# 1. 读取Word文档，按段落分割为LangChain文档对象
def load_word_to_docs(docx_path: str) -> List[LCDocument]:
    doc = Document(docx_path)
    docs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            docs.append(LCDocument(page_content=text))
    return docs

# 2. 构建知识库（RAG向量库）
def build_knowledge_base(docs: List[LCDocument]):
    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    splits = splitter.split_documents(docs)
    embeddings = DashScopeEmbeddings(dashscope_api_key="sk-2ccd6eee4dc04773add239ab18db4a8f")
    vectorstore = FAISS.from_documents(splits, embeddings)
    return vectorstore

# 3. 用LLM自动生成演绎大纲
def generate_narrative_outline(docs: List[LCDocument]) -> List[str]:
    context = "\n".join([d.page_content for d in docs[:10]])  # 取前10段做大纲
    prompt = f"请根据以下内容，列出讲解大纲（每一行一个知识点）：\n{context}\n\n大纲："
    result = llm.invoke(prompt)
    outline = [line.strip() for line in result.content.splitlines() if line.strip()]
    return outline

# 4. 用LLM+RAG生成某个知识点的图文讲解
def generate_narrative_step(topic: str, knowledge_base: Any) -> str:
    retriever = knowledge_base.as_retriever(search_kwargs={"k": 3})
    docs = retriever.invoke(topic)
    context = "\n".join([doc.page_content for doc in docs])
    prompt = f"请用通俗易懂的语言，结合以下知识，图文并茂地讲解：\n主题：{topic}\n相关内容：{context}\n\n讲解："
    result = llm.invoke(prompt)
    return result.content

# 5. RAG问答
def answer_question_with_rag(question: str, knowledge_base: Any) -> str:
    retriever = knowledge_base.as_retriever(search_kwargs={"k": 3})
    docs = retriever.invoke(question)
    context = "\n".join([doc.page_content for doc in docs])
    
    # 改进的提示词，让LLM知道这是关于课程的问题
    prompt = f"""你是一个专业的课程助手，正在回答关于"面向边缘场景动态资源的多出口神经网络设计方法和装置"这门课程的问题。

请结合以下课程内容，用友好、易懂的语言回答用户问题。如果问题比较宽泛，请结合课程内容给出具体的回答。

课程内容：
{context}

用户问题：{question}

请回答："""
    
    result = llm.invoke(prompt)
    return result.content

# 6. 生成测试题目
def generate_test_questions(learned_topics: List[str], knowledge_base: Any) -> List[dict]:
    """
    基于已学习的主题生成选择题
    参数：
        learned_topics: 已学习的主题列表
        knowledge_base: RAG知识库
    返回：
        测试题目列表
    """
    if not learned_topics:
        return []
    
    # 检索相关内容
    topics_text = "、".join(learned_topics)
    retriever = knowledge_base.as_retriever(search_kwargs={"k": 5})
    docs = retriever.invoke(f"关于{topics_text}的内容")
    context = "\n".join([doc.page_content for doc in docs])
    
    prompt = f"""基于以下学习内容，生成3道选择题来测试学生的掌握情况。

学习主题：{topics_text}

相关内容：
{context}

请按照以下格式生成题目：

题目1：
问题：[具体问题]
A. [选项A]
B. [选项B]
C. [选项C]
D. [选项D]
正确答案：[A/B/C/D]
解释：[为什么这个答案是正确的]

题目2：
[同样格式]

题目3：
[同样格式]

要求：
1. 题目要针对刚学过的内容
2. 难度适中，既不太简单也不太难
3. 选项要有一定的迷惑性
4. 解释要清晰易懂
"""
    
    try:
        result = llm.invoke(prompt)
        questions = parse_test_questions(result.content)
        return questions if questions else get_default_questions()
    except Exception as e:
        print(f"[DEBUG] 生成测试题目失败: {e}")
        return get_default_questions()

def parse_test_questions(llm_output: str) -> List[dict]:
    """
    解析LLM生成的测试题目文本
    """
    questions = []
    lines = llm_output.strip().split('\n')
    current_question = {}
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('题目') and '：' in line:
            if current_question:
                questions.append(current_question)
            current_question = {}
        elif line.startswith('问题：'):
            current_question['question'] = line.replace('问题：', '').strip()
            current_question['options'] = []
        elif line.startswith(('A.', 'B.', 'C.', 'D.')):
            if 'options' in current_question:
                current_question['options'].append(line)
        elif line.startswith('正确答案：'):
            answer = line.replace('正确答案：', '').strip()
            current_question['correct'] = answer
        elif line.startswith('解释：'):
            current_question['explanation'] = line.replace('解释：', '').strip()
    
    # 添加最后一个题目
    if current_question and len(current_question.get('options', [])) == 4:
        questions.append(current_question)
    
    return questions

def get_default_questions() -> List[dict]:
    """
    获取默认测试题目（当LLM生成失败时使用）
    """
    return [
        {
            "question": "关于神经网络的基本概念，以下哪个说法是正确的？",
            "options": [
                "A. 神经网络只能用于图像识别",
                "B. 神经网络是模拟人脑神经元工作原理的计算模型",
                "C. 神经网络不需要训练数据",
                "D. 神经网络只有一层结构"
            ],
            "correct": "B",
            "explanation": "神经网络是受生物神经网络启发的计算模型，通过模拟神经元之间的连接和信息传递来处理信息。"
        },
        {
            "question": "在边缘计算场景中，多出口神经网络的主要优势是什么？",
            "options": [
                "A. 增加网络复杂度",
                "B. 提高计算资源消耗",
                "C. 根据资源情况动态选择合适的输出",
                "D. 只能处理简单任务"
            ],
            "correct": "C",
            "explanation": "多出口神经网络可以根据边缘设备的计算资源和精度要求，动态选择合适的输出点，实现资源和性能的平衡。"
        },
        {
            "question": "动态资源分配在神经网络中的作用是什么？",
            "options": [
                "A. 固定使用最大资源",
                "B. 根据任务需求和设备能力灵活调整资源使用",
                "C. 总是使用最少资源",
                "D. 不考虑设备限制"
            ],
            "correct": "B",
            "explanation": "动态资源分配能够根据具体任务的复杂度和设备的计算能力，智能地分配和调整计算资源，提高整体效率。"
        }
    ]

# 7. 评估测试答案
def evaluate_test_answers(questions: List[dict], user_answers: List[str]) -> tuple:
    """
    评估用户的测试答案
    参数：
        questions: 测试题目列表
        user_answers: 用户答案列表
    返回：
        (得分, 详细结果列表)
    """
    if len(questions) != len(user_answers):
        return 0, []
    
    score = 0
    detailed_results = []
    
    for i, (question, user_answer) in enumerate(zip(questions, user_answers)):
        correct_answer = question.get('correct', 'A')
        is_correct = user_answer.upper() == correct_answer.upper()
        
        if is_correct:
            score += 1
        
        detailed_results.append((
            question,
            user_answer,
            is_correct,
            question.get('explanation', '暂无解释')
        ))
    
    return score, detailed_results

def generate_intro(docs: list, outline: list) -> str:
    """用LLM生成课程摘要+欢迎语+引导语"""
    context = "\n".join([d.page_content for d in docs[:10]])
    outline_str = "\n".join(outline)
    prompt = (
        f"请根据以下内容，生成一段课程开场白，包含：\n"
        f"1. 对文档知识的核心内容摘要\n"
        f"2. 对学习者的欢迎和鼓励\n"
        f"3. 简要展示教学大纲\n"
        f"4. 以‘准备好和我一起开始学习了吗？’结尾\n\n"
        f"文档内容：\n{context}\n\n大纲：\n{outline_str}\n\n开场白："
    )
    result = llm.invoke(prompt)
    return result.content

# 8. 智能分析测验结果和历史疑问
def analyze_test_results_and_questions(state: dict) -> str:
    """
    分析测验结果和历史疑问，生成针对性的学习建议
    参数：
        state: 包含测验结果和消息历史的状态
    返回：
        分析结果字符串，用于指导下一步教学
    """
    test_questions = state.get("test_questions", [])
    test_answers = state.get("test_answers", [])
    test_score = state.get("test_score", 0)
    messages = state.get("messages", [])
    
    # 分析测验结果
    test_analysis = ""
    if test_questions and test_answers:
        wrong_questions = []
        for i, (question, user_answer) in enumerate(zip(test_questions, test_answers)):
            correct_answer = question.get('correct', 'A')
            if user_answer.upper() != correct_answer.upper():
                wrong_questions.append({
                    'question': question.get('question', ''),
                    'user_answer': user_answer,
                    'correct_answer': correct_answer,
                    'explanation': question.get('explanation', '')
                })
        
        if wrong_questions:
            test_analysis = f"测验得分：{test_score}/{len(test_questions)}。"
            test_analysis += f"错误题目涉及：" + "、".join([q['question'][:20] + "..." for q in wrong_questions[:2]])
        else:
            test_analysis = f"测验表现优秀，得分：{test_score}/{len(test_questions)}。"
    
    # 分析历史疑问
    recent_questions = []
    for msg in reversed(messages[-10:]):  # 分析最近10条消息
        role = msg.get("role") if isinstance(msg, dict) else getattr(msg, "role", None)
        content = msg.get("content") if isinstance(msg, dict) else getattr(msg, "content", "")
        if role == "user" and content.strip() and content.lower() not in {
            "c", "continue", "继续", "开始", "好的", "ok", "yes", "测验", "test", "测试", "回顾", "review", "back"
        }:
            recent_questions.append(content)
            if len(recent_questions) >= 3:  # 最多分析3个最近的问题
                break
    
    questions_analysis = ""
    if recent_questions:
        questions_analysis = f"最近关注的问题包括：" + "、".join([q[:15] + "..." for q in recent_questions])
    
    # 使用LLM综合分析
    analysis_prompt = f"""请分析学生的学习情况并给出教学建议：

测验情况：{test_analysis if test_analysis else "暂无测验数据"}

最近疑问：{questions_analysis if questions_analysis else "暂无特别疑问"}

请简要分析学生可能的薄弱点，并给出1-2个关键词作为下一步教学的重点关注方向。
格式：重点关注：[关键词1]、[关键词2]
"""
    
    try:
        result = llm.invoke(analysis_prompt)
        analysis_result = result.content.strip()
        
        # 提取关键词
        if "重点关注：" in analysis_result:
            keywords = analysis_result.split("重点关注：")[1].strip()
            return keywords
        else:
            return analysis_result[:50]  # 返回前50个字符作为关键词
    except Exception as e:
        print(f"[DEBUG] 智能分析失败: {e}")
        # 简单的备用分析
        if test_analysis and "错误" in test_analysis:
            return "基础概念理解、实际应用"
        elif questions_analysis:
            return "深入理解、实践应用"
        else:
            return "全面掌握、系统理解"